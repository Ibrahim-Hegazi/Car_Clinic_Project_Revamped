# Library Name: python-docx
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# === Create Document ===
doc = Document()

# Title
doc.add_heading("Reddit LLM Pipeline Optimization Report", level=1)

# Problem Statement
doc.add_heading("🛠 Problem Statement", level=2)
problem_text = (
    "You have a daily data pipeline that collects approximately 700 Reddit posts related to automotive issues. "
    "These posts are passed to an LLM cleaning phase that extracts structured problem–solution pairs in JSON format. "
    "The LLM is run locally via Ollama, hosted on a GitHub-connected device without access to a GPU.\n\n"
    "Currently, the process runs serially and takes over 30 hours to complete a single day’s batch, "
    "rendering the system unusable for real-time or near-real-time production deployment."
)
doc.add_paragraph(problem_text)

# Root Causes
doc.add_heading("🔍 Root Causes", level=2)
root_causes = [
    "Serial Processing: Posts are sent to the LLM one by one, without any batching or parallelization.",
    "No GPU Acceleration: The local LLM runs on CPU, significantly slowing down token generation and response time.",
    "No External API: Cost constraints prevent the use of paid LLM APIs like OpenAI or Anthropic.",
    "Limited Hosting Options: The system avoids cloud GPU platforms due to budget limits.",
    "Captchas/Bot Blockers: Attempting automation via unofficial browser automation could face anti-bot protections (e.g., Cloudflare, CAPTCHA).",
    "Resource Bottlenecks: Limited threads, memory, or CPU on the hosting machine may further throttle performance."
]
for cause in root_causes:
    doc.add_paragraph(f"• {cause}", style='List Bullet')

# Solution Options Table
doc.add_heading("🧩 Solution Options Comparison Table", level=2)

headers = [
    "Solution Option", "💰 Monetary Cost", "🖥️ Hardware Requirements", "🕒 Dev Time",
    "🔧 Maintenance Overhead", "✅ Feasibility", "✅ Pros", "⚠️ Cons", "🔍 Notes / Hidden Risks"
]

data = [
    ["1. Use Official APIs", "$20–$100+/mo", "No hardware needed", "Low", "Low", "High",
     "Stable, reliable", "Ongoing cost in USD", "Budget constraints if prices rise"],

    ["2. Lightweight Queue System + Micro-batching", "$0", "Existing machine", "Medium", "Medium", "High",
     "Modular and scalable", "Still bound by CPU-only limits", "Needs smart batching + error handling"],

    ["3. Deploy to a More Powerful Local Machine", "$200–$1000 one-time", "New device (more RAM/CPU or GPU)", "Low",
     "Low", "Medium", "Fastest local option", "Upfront cost", "Hardware must be well configured for Ollama"],

    ["4. Browser Automation for Web LLMs (Playwright)", "$0–$20/mo (for proxies)", "No GPU needed", "High",
     "High", "Medium", "Can leverage powerful LLMs for free", "Captchas, bot blocking",
     "Proxy rotation, potential legal risk"],

    ["5. Use Reverse-Engineered APIs (⚠️ Risky)", "$0", "Internet connection only", "Medium", "High", "Low",
     "Access to high-quality models for free", "High ethical, legal, and stability risks", "May get blocked anytime"],

    ["6. Hybrid Model Design (Local + Remote)", "$0–$10/mo", "Small local model + Ollama", "High", "Medium", "Medium",
     "Offload easy cases, optimize cost", "Complex to implement", "Needs intelligent routing between models"],

    ["7. Parallel Local Execution with Ollama", "$0", "Multi-core CPU", "Medium", "Low", "High",
     "Fast improvement without new hardware", "Limited by CPU speed",
     "Needs careful multiprocessing to avoid overloading"],
]

# Create the table
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'  # Adds lines

# Set header row
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h

# Add data rows
for row in data:
    row_cells = table.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = item

# Final Recommendation
doc.add_heading("📌 Conclusion", level=2)
conclusion = (
    "After evaluating all the above options, **Solution 1: Use Official APIs** emerges as the most feasible for this project. "
    "It offers a stable and scalable pathway forward with minimal hardware dependencies and development time. "
    "Although it incurs a monetary cost, the reliability and speed gains far outweigh the expenses, especially if real-time performance is critical."
)
doc.add_paragraph(conclusion)

# Save
output_path = "Reddit_LLM_Pipeline_Optimization_Report.docx"
doc.save(output_path)

print(f"✅ Document saved to: {output_path}")
